from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import io
import re

app = FastAPI()

# Allow requests from the Next.js dev server and production origin.
# Update ALLOWED_ORIGINS if your Next.js app runs on a different port.
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:3001",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["POST"],
    allow_headers=["Content-Type"],
)

# ─── Pydantic models matching TemplateData shape ──────────────────────────────
# These mirror the TypeScript types in types/pageTypes.ts so FastAPI can
# validate the incoming JSON and provide clear error messages on bad input.

class TableData(BaseModel):
    id: str
    rows: int
    cols: int
    data: list[list[str]]

class SectionNode(BaseModel):
    id: str
    number: str
    title: str
    content: str
    # Support both the old single 'locked' field and the new granular lock fields.
    # If 'locked' is absent (new schema), derive it from lockEdit.
    locked: Optional[bool] = None
    lockEdit: Optional[bool] = None
    lockDelete: Optional[bool] = None
    lockAddTable: Optional[bool] = None
    lockAddSections: Optional[bool] = None
    tables: Optional[list[TableData]] = []
    children: Optional[list["SectionNode"]] = []

    @property
    def is_locked(self) -> bool:
        """True if content is locked — works with both old and new schema."""
        if self.locked is not None:
            return self.locked
        return bool(self.lockEdit)

# Required for self-referential Pydantic model (children: list[SectionNode])
SectionNode.model_rebuild()

class HeaderFooterData(BaseModel):
    headerLeft: str
    headerCenter: str
    headerRight: str
    footerLeft: str
    footerCenter: str
    footerRight: str
    showPageNumbers: bool
    pageNumberPosition: str

class TemplateField(BaseModel):
    id: str
    label: str
    type: str
    defaultValue: Optional[str] = ""
    placeholder: Optional[str] = ""
    required: Optional[bool] = False

class TemplateData(BaseModel):
    documentName: str
    fields: list[TemplateField]
    headerFooter: HeaderFooterData
    sections: list[SectionNode]


# ─── Token resolver ───────────────────────────────────────────────────────────
# Replaces {{field_id}} tokens in content strings with their defaultValue.
# The engineer page merges fieldValues into defaultValue before sending here,
# so defaultValue holds the final answer the engineer typed.
def resolve_tokens(content: str, field_map: dict[str, str]) -> str:
    def replace(match):
        field_id = match.group(1)
        return field_map.get(field_id, f"[{field_id}]")
    return re.sub(r"\{\{([^}]+)\}\}", replace, content)


# ─── Section renderer ─────────────────────────────────────────────────────────
# Recursively adds sections to the document with correct heading levels.
# Depth 0 = Heading 1, depth 1 = Heading 2, depth 2 = Heading 3.
def add_sections(doc: Document, sections: list[SectionNode], field_map: dict[str, str], depth: int = 0):
    heading_level = min(depth + 1, 3)
    BODY_INDENT = Inches(0.3)

    for section in sections:
        # Section heading — number + title
        heading_text = f"{section.number}  {section.title}"
        heading = doc.add_heading("", level=heading_level)
        run = heading.add_run(heading_text)

        style_run(
            run,
            font_size=18 if heading_level == 1 else 14,
            bold=True
        )

        if heading_level > 1:
            heading.paragraph_format.left_indent = Inches(0.3 * (heading_level - 1))
        
        

        # Section body — resolve any {{tokens}} to their filled values
        INDENT = Inches(0.3)

        if section.content.strip():
            resolved = resolve_tokens(section.content, field_map)

            para = doc.add_paragraph()

            base_indent = 0.3  # inches

            para.paragraph_format.left_indent = Inches(base_indent * (heading_level - 1))
            para.paragraph_format.first_line_indent = Inches(0.3)

            run = para.add_run(resolved)
            style_run(run, font_size=12)

        # Tables
        if section.tables:
            for table_data in section.tables:
                if table_data.rows > 0 and table_data.cols > 0:
                    table = doc.add_table(rows=table_data.rows, cols=table_data.cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(table_data.data):
                        for ci, cell_value in enumerate(row):
                            if ri < len(table.rows) and ci < len(table.rows[ri].cells):
                                table.rows[ri].cells[ci].text = cell_value
                    doc.add_paragraph()  # spacing after table

        # Recurse into children
        if section.children:
            add_sections(doc, section.children, field_map, depth + 1)


# ─── Main export endpoint ─────────────────────────────────────────────────────
@app.post("/generate")
async def generate_docx(template: TemplateData):
    try:
        doc = Document()

        

        for i in range(1, 4):
            style = doc.styles[f"Heading {i}"]
            rPr = style._element.get_or_add_rPr()

            # Remove theme color if it exists
            for elem in rPr.findall(qn('w:color')):
                rPr.remove(elem)

            # Force black color
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "000000")
            rPr.append(color)

        # Build a field map of id -> resolved value for token replacement.
        field_map = {
            f.id: (f.defaultValue or f.placeholder or f"[{f.label}]")
            for f in template.fields
        }

        # ── Cover page ────────────────────────────────────────────────────────
        title_val = field_map.get("field_cover_title", "STATEMENT OF WORK")
        client_val = field_map.get("field_cover_client_name", "—")
        bldg_val = field_map.get("field_cover_building", "")
        loc_val = field_map.get("field_cover_location", "—")
        prepared_val = field_map.get("field_cover_prepared_by", "—")
        dept_val = field_map.get("field_cover_department", "—")
        date_val = field_map.get("field_cover_date", "—")
        version_val = field_map.get("field_cover_version", "1.0")
        proj_num_val = field_map.get("field_project_number", "—")
        conf_val = field_map.get("field_cover_confidentiality", "Confidential")

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title_val.upper())
        title_run.bold = True
        title_run.font.size = Pt(24)

        doc.add_paragraph()

        for_para = doc.add_paragraph()
        for_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for_run = for_para.add_run("FOR")
        for_run.bold = True
        for_run.font.size = Pt(18)

        client_para = doc.add_paragraph()
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client_para.add_run(client_val.upper() if client_val != "—" else "—")
        client_run.bold = True
        client_run.font.size = Pt(18)

        if bldg_val and bldg_val != "—" and bldg_val != "[Building]":
            bldg_para = doc.add_paragraph()
            bldg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bldg_run = bldg_para.add_run(f"BUILDING {bldg_val}")
            bldg_run.bold = True
            bldg_run.font.size = Pt(14)

        doc.add_paragraph()
        doc.add_paragraph()

        details = [
            loc_val,
            "",
            "Prepared by",
            prepared_val,
            dept_val,
            "",
            f"Date: {date_val}",
            f"Version: {version_val}",
            f"Project Number: {proj_num_val}",
            f"Classification: {conf_val}",
        ]
        for line in details:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── End cover page with a section break that carries border + vAlign ──
        # Embedding a w:sectPr inside a paragraph's w:pPr ends Section 1 here.
        # Everything before this paragraph is in Section 1 (cover page only).
        # Everything after belongs to Section 2 (the rest of the document).
        # This isolates the border and vertical centering to the cover page only.
        cover_break_para = doc.add_paragraph()
        cover_pPr = OxmlElement("w:pPr")

        cover_sectPr = OxmlElement("w:sectPr")

        # Vertical centering — centers cover content between top and bottom margins
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        cover_sectPr.append(vAlign)

        # Page border — single black 3pt border on all four sides, cover page only
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")
        for side in ("w:top", "w:left", "w:bottom", "w:right"):
            b = OxmlElement(side)
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "24")
            b.set(qn("w:space"), "24")
            b.set(qn("w:color"), "000000")
            pgBorders.append(b)
        cover_sectPr.append(pgBorders)

        cover_pPr.append(cover_sectPr)
        cover_break_para._p.append(cover_pPr)

        # ── Table of Contents with real Word TOC field ────────────────────────
        # Inserts a proper TOC field that Word recognizes and can update.
        # When the engineer opens the document they press Ctrl+A then F9
        # to update all fields which generates the real TOC automatically.
        doc.add_heading("Table of Contents", level=1)

        # The TOC field uses a fldChar/instrText sequence that Word reads
        # and builds into a real TOC from the document headings.
        # \o "1-3" includes heading levels 1 through 3
        # \h makes entries into hyperlinks
        # \z hides tab leader in web layout
        # \u uses the applied paragraph outline level
        # dirty="true" tells Word to update the field when the document opens
        toc_para = doc.add_paragraph()

        run1 = toc_para.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        fldChar_begin.set(qn("w:dirty"), "true")
        run1._r.append(fldChar_begin)

        run2 = toc_para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)

        run3 = toc_para.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

        note = doc.add_paragraph("Press Ctrl+A then F9 to generate the Table of Contents.")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_page_break()

        # ── Document sections ─────────────────────────────────────────────────
        add_sections(doc, template.sections, field_map, depth=0)

        # ── Header and footer (Section 2 — all pages after cover) ─────────────
        # Tab stops position left/center/right text in a single paragraph
        # without tables so no visible borders appear.
        hf = template.headerFooter
        # 6.5in usable width in twips (1in = 1440 twips)
        PAGE_WIDTH_TWIPS = 9360

        def make_page_number_field() -> list:
            """Returns XML run elements forming a real Word PAGE field."""
            r1 = OxmlElement("w:r")
            fc1 = OxmlElement("w:fldChar")
            fc1.set(qn("w:fldCharType"), "begin")
            r1.append(fc1)

            r2 = OxmlElement("w:r")
            instr = OxmlElement("w:instrText")
            instr.text = " PAGE "
            instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r2.append(instr)

            r3 = OxmlElement("w:r")
            fc3 = OxmlElement("w:fldChar")
            fc3.set(qn("w:fldCharType"), "end")
            r3.append(fc3)

            return [r1, r2, r3]

        def make_run(text: str, font_size: int = 9) -> OxmlElement:
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(font_size * 2))
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            return r

        def make_tab_run() -> OxmlElement:
            r = OxmlElement("w:r")
            r.append(OxmlElement("w:tab"))
            return r

        def add_hf_paragraph(parent, left_text: str, center_text: str, right_text: str):
            """
            Single paragraph with center and right tab stops.
            Handles {PAGE} token by inserting a real Word PAGE field.
            """
            for para in parent.paragraphs:
                para._p.getparent().remove(para._p)

            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")

            tab_c = OxmlElement("w:tab")
            tab_c.set(qn("w:val"), "center")
            tab_c.set(qn("w:pos"), str(PAGE_WIDTH_TWIPS // 2))
            tabs.append(tab_c)

            tab_r = OxmlElement("w:tab")
            tab_r.set(qn("w:val"), "right")
            tab_r.set(qn("w:pos"), str(PAGE_WIDTH_TWIPS))
            tabs.append(tab_r)

            pPr.append(tabs)
            p.append(pPr)

            def append_text_or_field(zone_text: str):
                """Splits text on {PAGE} and inserts a real field for that token."""
                if not zone_text:
                    return
                parts = zone_text.split("{PAGE}")
                for i, part in enumerate(parts):
                    if part:
                        p.append(make_run(part))
                    if i < len(parts) - 1:
                        for field_run in make_page_number_field():
                            p.append(field_run)

            append_text_or_field(left_text)
            p.append(make_tab_run())
            append_text_or_field(center_text)
            p.append(make_tab_run())
            append_text_or_field(right_text)

            parent._element.append(p)

        # Section 2 is doc.sections[0] in python-docx when a section break exists
        # doc.sections[-1] always refers to the last (main) section
        main_section = doc.sections[-1]

        add_hf_paragraph(
            main_section.header,
            resolve_tokens(hf.headerLeft.split("\n")[0] if hf.headerLeft else "", field_map),
            resolve_tokens(hf.headerCenter, field_map),
            resolve_tokens(hf.headerRight, field_map),
        )
        add_hf_paragraph(
            main_section.footer,
            resolve_tokens(hf.footerLeft, field_map),
            resolve_tokens(hf.footerCenter, field_map),
            resolve_tokens(hf.footerRight, field_map),
        )

        # ── Serialize and return ──────────────────────────────────────────────
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = template.documentName.replace(" ", "-").lower()
        headers = {
            "Content-Disposition": f'attachment; filename="{filename}.docx"'
        }

        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Health check ─────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {"status": "ok"}


# Style Changes for parity with SOW-Editor
def style_run(run, font_name="Calibri", font_size=11, color=(0, 0, 0), bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # Required for Word to fully respect font name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.color.rgb = RGBColor(*color)